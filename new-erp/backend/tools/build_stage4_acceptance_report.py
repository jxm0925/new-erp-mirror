from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\new-erp\new-erp\backend")
DOCS_ROOT = Path(r"D:\codex-introduce\new_erp\docs")
OUT_DIR = DOCS_ROOT / "order-to-work-order" / "step04-sku-item-matching"
CHECK_DIR = DOCS_ROOT / "ui-check" / "order-to-work-order" / "step04-sku-item-matching"
OUT_FILE = OUT_DIR / "第四阶段_SKU-Item默认关系最终验收报告_20260725.docx"

GREEN = "07883D"
DARK_GREEN = "075E2B"
LIGHT_GREEN = "EAF7EF"
BLUE = "245B86"
DARK = "17212B"
GRAY = "667085"
LIGHT_GRAY = "F2F4F7"
ORANGE = "B54708"
LIGHT_ORANGE = "FFF4E5"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, bold=False, color=DARK, size=9.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)


def style_table(table, widths, header=True):
    set_table_geometry(table, widths)
    for row_idx, row in enumerate(table.rows):
        if header and row_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
            for cell in row.cells:
                shade(cell, LIGHT_GRAY)
                for run in cell.paragraphs[0].runs:
                    set_run_font(run, size=9.2, bold=True, color=DARK)
        else:
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    set_run_font(run, size=9.2, color=DARK)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True, color=DARK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=10.5, color=DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=DARK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.2, color=DARK)


def add_status_box(doc):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, LIGHT_GREEN)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("验收结论：通过（第四阶段功能开发与技术门禁完成）")
    set_run_font(r, size=12, bold=True, color=DARK_GREEN)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(
        "当前正式数据检查结果为 12 个实物 SKU：1 个关系正常、11 个缺少默认 Item。"
        "缺失项属于待维护的业务主数据，不允许系统猜测绑定，不构成代码缺陷。"
    )
    set_run_font(r2, size=9.8, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_result_table(doc):
    rows = [
        ("验收项", "结果", "结论摘要"),
        ("关系列表", "通过", "真实接口分页；筛选、状态、设置/更换、历史入口完整；表格无横向裁切。"),
        ("设置/更换默认 Item", "通过", "只允许实物 SKU；候选 Item 真实分页；禁用/同 Item/服务/无需发货均阻止保存。"),
        ("关系详情与历史", "通过", "当前关系、历史快照、操作日志同页展示；异常信息按状态出现。"),
        ("完整性检查", "通过", "只读检查、状态筛选、重新检查、修复入口及分页有效，不自动改业务数据。"),
        ("状态机与事务", "通过", "唯一有效默认关系、旧关系失效、历史/日志落库、失败整体回滚。"),
        ("RBAC 按钮权限", "通过", "查看、检查、设置/更换、修复接口分别校验，越权返回 403。"),
        ("锁定设计图对齐", "通过", "四张已通过设计图未重新设计，仅按后续明确指令做局部业务修正。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    for idx, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[idx], value, bold=True)
    for item in rows[1:]:
        cells = table.add_row().cells
        set_cell_text(cells[0], item[0])
        set_cell_text(cells[1], item[1], bold=True, color=GREEN)
        set_cell_text(cells[2], item[2])
    style_table(table, [2100, 1100, 6160])


def add_test_table(doc):
    rows = [
        ("验证层级", "执行结果", "证据"),
        ("后端全量自动化", "20 项 / 84 断言全部通过", "SQLite 独立测试库；无失败、无跳过。"),
        ("第四阶段专项", "9 项 / 37 断言全部通过", "覆盖状态、事务、历史、日志、权限、真实分页和只读检查。"),
        ("前端生产构建", "通过", "Vue CLI production build 成功；仅保留既有 bundle 体积警告。"),
        ("浏览器点击验收", "通过", "管理员真实登录，逐页点击筛选、分页、详情、设置页、完整性检查。"),
        ("浏览器控制台", "0 个 error / 0 个 warning", "第四阶段验收路径未发现前端运行时错误。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    for idx, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[idx], value, bold=True)
    for item in rows[1:]:
        cells = table.add_row().cells
        set_cell_text(cells[0], item[0])
        set_cell_text(cells[1], item[1], bold=True, color=GREEN)
        set_cell_text(cells[2], item[2])
    style_table(table, [2050, 2600, 4710])


def add_browser_table(doc):
    rows = [
        ("页面/动作", "实际结果"),
        ("关系列表切换第 2 页", "接口返回真实第 2 页，页面显示 2 条记录；总数 12、每页 10。"),
        ("打开 SKU000002 关系详情", "展示当前 Item、关系历史和操作日志，业务字段与数据库一致。"),
        ("点击“更换默认 Item”", "进入独立页面；候选 Item 共 13 条、2 页，翻页有效。"),
        ("选择真实候选 Item", "右侧预览显示真实编码、名称、类型和单位；测试未提交，不污染数据。"),
        ("选择原因“其他”后直接保存", "页面提示必须填写备注并阻止提交，URL 保持不变。"),
        ("完整性检查切换“正常/全部”", "正常 1 条、全部 12 条，与接口统计一致。"),
        ("点击“重新检查”", "结果仍为正常 1、缺失 11、异常 0、无需 Item 0；检查过程无写入。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    for idx, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[idx], value, bold=True)
    for item in rows[1:]:
        cells = table.add_row().cells
        set_cell_text(cells[0], item[0])
        set_cell_text(cells[1], item[1])
    style_table(table, [3300, 6060])


def add_artifact_table(doc):
    rows = [
        ("类别", "位置"),
        ("锁定设计图", r"D:\codex-introduce\new_erp\docs\ui-reference\order-to-work-order\step04-sku-item-matching\approved"),
        ("验收截图", r"D:\codex-introduce\new_erp\docs\ui-check\order-to-work-order\step04-sku-item-matching"),
        ("后端专项测试", r"tests\Feature\Erp\SkuItemDefaultRelationServiceTest.php"),
        ("验收报告", r"D:\codex-introduce\new_erp\docs\order-to-work-order\step04-sku-item-matching"),
    ]
    table = doc.add_table(rows=1, cols=2)
    for idx, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[idx], value, bold=True)
    for item in rows[1:]:
        cells = table.add_row().cells
        set_cell_text(cells[0], item[0])
        set_cell_text(cells[1], item[1], size=8.6)
    style_table(table, [2100, 7260])


def add_comparison(doc, title, filename, note):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=BLUE)
    path = CHECK_DIR / filename
    doc.add_picture(str(path), width=Inches(6.38))
    cp = doc.paragraphs[-1]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(3)
    np = doc.add_paragraph()
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np.paragraph_format.space_after = Pt(5)
    nr = np.add_run(note)
    set_run_font(nr, size=8.4, color=GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.space_after = Pt(6)

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_GREEN, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    set_cell_width(table.cell(0, 0), 4680)
    set_cell_width(table.cell(0, 1), 4680)
    p1 = table.cell(0, 0).paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("ERP 重构｜第四阶段验收")
    set_run_font(r1, size=8.5, bold=True, color=GRAY)
    p2 = table.cell(0, 1).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("SKU–Item 默认关系")
    set_run_font(r2, size=8.5, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("山东简探 ERP  |  2026-07-25  |  第 ")
    set_run_font(fr, size=8, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)
    fr2 = fp.add_run(" 页")
    set_run_font(fr2, size=8, color=GRAY)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("第四阶段最终验收报告")
    set_run_font(r, size=23, bold=True, color=DARK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("SKU–Item 默认关系管理")
    set_run_font(sr, size=13, color=GRAY)

    mast = doc.add_table(rows=5, cols=2)
    metadata = [
        ("项目", "ERP 系统重构"),
        ("阶段", "订单到生产工单全流程 · 第四项"),
        ("验收日期", "2026-07-25"),
        ("验收环境", "Laravel API + Vue 管理端 + MySQL 正式开发库"),
        ("最终结论", "通过（功能开发与技术门禁完成）"),
    ]
    for i, (label, value) in enumerate(metadata):
        set_cell_text(mast.cell(i, 0), label, bold=True, color=GRAY)
        set_cell_text(mast.cell(i, 1), value, bold=i == 4, color=GREEN if i == 4 else DARK)
    set_table_geometry(mast, [1800, 7560])
    for row in mast.rows:
        shade(row.cells[0], LIGHT_GRAY)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(10)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GREEN)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    add_status_box(doc)
    add_heading(doc, "1. 验收范围与结论", 1)
    add_body(
        doc,
        "本次验收覆盖 SKU–Item 默认关系的关系列表、设置/更换默认 Item、关系详情与历史、关系完整性检查四个锁定页面，"
        "以及对应接口、状态机、事务、历史快照、操作日志、RBAC 和前后端分页。"
    )
    add_result_table(doc)

    add_heading(doc, "2. 已固化的业务规则", 1)
    add_bullets(
        doc,
        [
            "只有“实物”订单行类型的 SKU 可以设置默认 Item；服务和无需发货 SKU 不允许绑定 Item。",
            "每个实物 SKU 在任一时刻最多只有一个启用且主关系；更换后旧关系立即失效并保留历史。",
            "候选 Item 必须启用；同一个 Item 不能重复设为当前默认 Item。",
            "设置或更换必须选择变更原因；选择“其他”时备注必填。",
            "完整性检查是只读审计，不自动修复、不猜测业务关系；修复动作必须由有权限的人员明确确认。",
            "历史销售订单继续使用下单时保存的 Item 快照，不受后续默认关系变更影响。",
        ],
    )

    add_heading(doc, "3. 自动化与构建验证", 1)
    add_test_table(doc)
    add_body(
        doc,
        "专项测试覆盖：首次设置、更换、禁用 Item、同 Item、服务/无需发货 SKU、“其他”无备注、重复关系修复、"
        "错误绑定解除、只读检查、日志失败事务回滚、按钮权限 403、列表及检查接口真实分页。"
    )

    add_heading(doc, "4. 用户视角页面点击验收", 1)
    add_browser_table(doc)
    add_body(
        doc,
        "测试过程中只选择真实候选 Item 做预览与校验，没有提交新的业务映射；浏览器控制台未出现 error 或 warning。"
    )

    add_heading(doc, "5. 当前正式数据完整性", 1)
    data_table = doc.add_table(rows=2, cols=5)
    headers = ("已检查", "正常", "缺少默认 Item", "关系异常", "无需 Item")
    values = ("12", "1", "11", "0", "0")
    for i, value in enumerate(headers):
        set_cell_text(data_table.cell(0, i), value, bold=True)
    for i, value in enumerate(values):
        color = GREEN if i == 1 else ORANGE if i == 2 else DARK
        set_cell_text(data_table.cell(1, i), value, bold=True, color=color, size=13)
    style_table(data_table, [1500, 1500, 2360, 2000, 2000])
    warning = doc.add_table(rows=1, cols=1)
    set_table_geometry(warning, [9360])
    shade(warning.cell(0, 0), LIGHT_ORANGE)
    set_cell_text(
        warning.cell(0, 0),
        "说明：11 个缺失关系必须依据真实物料主数据由业务人员逐项确认。系统不会根据名称相似度或示例数据自动绑定，"
        "因此第四阶段可以判定“功能完成”，但不能判定“全部业务主数据已配置完成”。",
        bold=True,
        color=ORANGE,
        size=9.6,
    )

    add_heading(doc, "6. 审计与历史数据说明", 1)
    add_body(
        doc,
        "现有历史中仍有一条操作人显示“系统”的旧记录，该记录形成于当前操作人归属修正之前。为保持审计真实性，本次未篡改历史归属。"
        "后续通过页面执行的设置、更换与修复操作会写入当前登录管理员及对应操作日志。"
    )

    add_heading(doc, "7. 设计图锁定与局部修正", 1)
    add_bullets(
        doc,
        [
            "四张已验收 ProductDesign 图继续作为唯一视觉基线，本次没有重新设计。",
            "真实系统壳层、真实数据数量与设计图示例数据不同，不视为视觉偏差。",
            "设置页增加候选 Item 分页，属于“所有列表前后端必须分页”的强制规则。",
            "详情页正常状态只显示 SKU 信息和当前默认 Item；异常卡片仅在确有异常时出现，遵循后续明确业务修正。",
            "表格列宽和操作列仅做防裁切修正，页面层级、操作入口和信息结构保持锁定设计。",
        ],
    )

    add_heading(doc, "8. 验收证据位置", 1)
    add_artifact_table(doc)

    add_heading(doc, "最终签署结论", 1)
    final_table = doc.add_table(rows=1, cols=1)
    set_table_geometry(final_table, [9360])
    shade(final_table.cell(0, 0), LIGHT_GREEN)
    set_cell_text(
        final_table.cell(0, 0),
        "第四阶段 SKU–Item 默认关系管理已完成开发、技术测试、生产构建、视觉核对和用户视角点击验收，"
        "满足进入下一项开发的技术条件。11 个缺失映射转为业务主数据维护清单，待获得真实映射依据后人工处理。",
        bold=True,
        color=DARK_GREEN,
        size=10.5,
    )

    add_heading(doc, "附录 A：锁定设计图与实装截图对比", 1)
    add_body(
        doc,
        "以下对比图左侧为已通过的锁定 ProductDesign，右侧为 2026-07-25 浏览器实装截图。"
        "比较重点为页面层级、控件位置、操作入口、表格结构和业务信息组织。"
    )
    add_comparison(
        doc,
        "A.1 关系列表",
        "sku-item-relation-list-comparison.png",
        "实装使用真实数据与真实分页；列表结构、筛选区、统计区、操作入口和说明栏保持一致。",
    )
    add_comparison(
        doc,
        "A.2 设置/更换默认 Item",
        "sku-item-set-primary-comparison.png",
        "实装保留独立页面和三栏信息结构，并加入真实候选 Item 分页与强校验。",
    )
    doc.add_page_break()
    add_comparison(
        doc,
        "A.3 关系详情与历史",
        "sku-item-relation-history-comparison.png",
        "实装按真实正常状态隐藏异常卡片；历史与操作日志结构保持锁定设计。",
    )
    add_comparison(
        doc,
        "A.4 关系完整性检查",
        "sku-item-integrity-check-comparison.png",
        "实装统计、筛选、检查口径、修复入口和分页均使用真实接口数据。",
    )

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build()
